"""Amostragem estratégica de documentos (cabeça/meio/cauda).

Em vez de jogar o documento inteiro no prompt (estoura contexto e custa caro
em planilhas de 2000+ linhas), extraímos AMOSTRAS representativas:

  · PDF   → primeiras 2 páginas + 1 do meio + última, texto por página.
  · XLSX  → por sheet: cabeçalho + ~25 primeiras linhas + ~10 últimas, +
            dimensões totais. Uma passada só, memória limitada (deque).
  · DOCX  → parágrafos + amostra de tabelas.
  · CSV   → primeiras ~40 linhas.
  · MD/TXT→ início do texto.

O agente recebe essas amostras + as DIMENSÕES TOTAIS e escreve o texto-mapa
descrevendo onde estão os dados — sabendo que viu só uma fatia, não o todo.

Tudo é defensivo: lib faltando, arquivo corrompido ou formato não suportado
viram `readable=False` + uma nota explicativa (o mapeamento segue, com
confiança baixa, ao invés de quebrar a fila).
"""

from __future__ import annotations

import io
import re
import unicodedata
from collections import deque
from dataclasses import dataclass, field
from typing import Any

# Limites de amostragem (bound de prompt/memória) ─────────────────────
_PDF_HEAD_PAGES = 2
_PDF_PER_PAGE_CHARS = 3500
_XLSX_HEAD_ROWS = 25
_XLSX_TAIL_ROWS = 10
_XLSX_MID_ROWS = 8  # janela do meio (fecha o ponto cego de blocos empilhados)
_XLSX_MAX_COLS = 30
_XLSX_CELL_CHARS = 60

# Workbook-motor (planilha multi-aba com índice/guia) ─────────────────
# Acima deste nº de abas, ativamos o modo "workbook": inventário de TODAS as abas primeiro, as
# abas-meta (índice/Guia) renderizadas quase inteiras e priorizadas, e um teto de amostra maior.
# Resolve o gargalo onde as 3 primeiras abas comem os 16K e as outras N ficam no escuro.
_XLSX_WORKBOOK_SHEETS = 6
_XLSX_META_HEAD_ROWS = 130  # abas-meta (Guia/MAPA): são o MAPA do doc → rende quase inteiras
# Hints de aba-meta (sem acento, minúsculo). Aba de conteúdo codificada (C.1, D.0…) é excluída
# antes, pra "C.14 Mapa da Obra" não virar falso-positivo de "mapa".
_XLSX_META_HINTS = ("guia", "instruc", "indice", "dicion", "mapa", "sumario", "legenda")
_CODE_PREFIX_RE = re.compile(r"^[A-Z]\s*[.\d]")  # 'C.1', 'D.0', 'B.4'… = aba de conteúdo
# "Grande" pra chunking = volume de CÉLULAS (linhas×colunas), não só linhas:
# uma sheet de 200 linhas × 600 colunas explode a saída tanto quanto 2000 linhas.
_XLSX_BIG_CELLS = 4000
_DOCX_CHARS = 9000
_TEXT_CHARS = 12000
_CSV_LINES = 40


@dataclass
class DocSample:
    """Resultado da amostragem · alimenta o prompt do mapeador."""

    text: str  # amostra textual (markdown) injetada no prompt
    fmt: str  # 'pdf' | 'xlsx' | 'docx' | 'csv' | 'md' | 'unknown'
    total_units: int  # páginas (pdf) | sheets (xlsx) | 0
    readable: bool  # False = não extraímos texto (escaneado/binário/corrompido)
    truncated: bool  # True = amostra cortada por limite
    structure: dict[str, Any] = field(default_factory=dict)  # mapa p/ persistência
    notes: list[str] = field(default_factory=list)  # avisos pro mapeador


def build_doc_samples(filename: str, data: bytes) -> DocSample:
    """Dispatcher por extensão. Nunca levanta exceção — degrada pra unknown."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    try:
        if ext == "pdf":
            return _sample_pdf(filename, data)
        if ext in ("xlsx", "xlsm"):
            return _sample_xlsx(filename, data)
        if ext == "xls":
            return _sample_xls(filename, data)
        if ext == "docx":
            return _sample_docx(filename, data)
        if ext == "csv":
            return _sample_text_lines(filename, data, fmt="csv")
        if ext in ("md", "markdown", "txt"):
            return _sample_text_chars(filename, data, fmt="md")
        if ext in ("doc", "mpp"):
            return _unsupported(
                filename,
                ext,
                f"Formato legado/binário .{ext} não tem parser direto aqui. "
                "Descreva pelo nome e marque que a extração precisará de conversão "
                "(ex: salvar como .docx, ou MS Project → export XML).",
            )
        return _unsupported(filename, ext or "?", "Extensão não reconhecida.")
    except Exception as e:  # noqa: BLE001 — robustez: amostragem nunca derruba a fila
        return DocSample(
            text=f"(falha ao amostrar '{filename}': {type(e).__name__}: {e})",
            fmt=ext or "unknown",
            total_units=0,
            readable=False,
            truncated=False,
            structure={"format": ext or "unknown", "totalUnits": 0, "error": str(e)[:300]},
            notes=[f"Erro de parsing: {type(e).__name__}. Arquivo pode estar corrompido."],
        )


# ── PDF ────────────────────────────────────────────────────────────
def _sample_pdf(filename: str, data: bytes) -> DocSample:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    n = len(reader.pages)
    if n == 0:
        return _unsupported(filename, "pdf", "PDF sem páginas.")

    # Índices amostrados: 2 primeiras + meio + última (sem repetir).
    idxs = sorted({0, 1, n // 2, n - 1} & set(range(n)))
    if n <= _PDF_HEAD_PAGES + 1:
        idxs = list(range(n))

    parts: list[str] = []
    any_text = False
    for i in idxs:
        try:
            raw = reader.pages[i].extract_text() or ""
        except Exception:  # noqa: BLE001
            raw = ""
        raw = raw.strip()
        if raw:
            any_text = True
        clipped = raw[:_PDF_PER_PAGE_CHARS]
        body = clipped if clipped else "(sem texto extraível nesta página)"
        parts.append(f"### Página {i + 1} de {n}\n{body}")

    notes: list[str] = []
    if not any_text:
        notes.append(
            "PDF sem texto extraível em nenhuma amostra — provavelmente "
            "escaneado/imagem. A extração precisará de OCR ou leitura por visão."
        )

    text, truncated = _clip("\n\n".join(parts))
    return DocSample(
        text=text,
        fmt="pdf",
        total_units=n,
        readable=any_text,
        truncated=truncated,
        structure={
            "format": "pdf",
            "totalUnits": n,
            "sampledPages": [i + 1 for i in idxs],
            "anomalies": [],
            "chunkingStrategy": (
                f"PDF com {n} páginas — extrair por blocos de páginas."
                if n > 20
                else None
            ),
            "sampledBy": "python-sampler@1.0.0",
        },
        notes=notes,
    )


# ── Workbook multi-aba: detecção de aba-meta + inventário ─────────────
def _strip_accents(s: str) -> str:
    return unicodedata.normalize("NFKD", s or "").encode("ascii", "ignore").decode("ascii")


def _is_meta_sheet(name: str) -> bool:
    """Aba de índice/guia/dicionário — o MAPA do workbook (vale priorizar quase inteira).
    Exclui abas de conteúdo codificadas (C.14 'Mapa da Obra' não é meta). Match por TOKEN
    (prefixo de palavra), não substring — senão 'aDICIONal' casaria 'dicion', 'maPAmento' casaria
    'mapa', etc."""
    raw = (name or "").strip()
    if _CODE_PREFIX_RE.match(raw):
        return False
    tokens = re.split(r"[^a-z0-9]+", _strip_accents(raw).lower())
    return any(tok.startswith(h) for tok in tokens for h in _XLSX_META_HINTS)


def _sheet_inventory(sheets_struct: list[dict[str, Any]]) -> str:
    """Bloco de INVENTÁRIO: TODA aba por nome + dimensões. Vai no topo da amostra (nunca cortado)
    → o mapeador enxerga as N abas mesmo quando o conteúdo de muitas não couber na amostra."""
    lines = [f"### INVENTÁRIO · {len(sheets_struct)} abas (nome · linhas×colunas)"]
    for s in sheets_struct:
        flag = " ⟵ META/índice (mapa do doc)" if s.get("isMeta") else ""
        lines.append(f'- "{s["name"]}" · {s["totalRows"]}×{s["totalCols"]}{flag}')
    return "\n".join(lines)


def _assemble_xlsx(
    blocks: list[tuple[bool, str]],
    sheets_struct: list[dict[str, Any]],
    workbook_mode: bool,
) -> tuple[str, bool, list[str]]:
    """Junta os blocos das sheets e corta no teto. No modo workbook: INVENTÁRIO no topo (nunca
    cortado) → abas-meta (índice/Guia) → abas de conteúdo, com teto MAIOR. Fora dele: ordem natural
    + teto global (comportamento legado, regressão-safe). Retorna (texto, truncado, notes)."""
    if not workbook_mode:
        text, truncated = _clip("\n\n".join(b for (_, b) in blocks))
        return text, truncated, []

    from config import MAX_SAMPLE_CHARS_WORKBOOK

    metas = [s["name"] for s in sheets_struct if s.get("isMeta")]
    inventory = _sheet_inventory(sheets_struct)
    meta_blocks = [b for (m, b) in blocks if m]
    content_blocks = [b for (m, b) in blocks if not m]
    ordered = [inventory] + meta_blocks + content_blocks
    text, truncated = _clip("\n\n".join(ordered), cap=MAX_SAMPLE_CHARS_WORKBOOK)

    notes = [
        f"Workbook multi-aba ({len(sheets_struct)} abas) — modo workbook-motor: inventário completo "
        "no topo, abas-meta priorizadas. "
        + (f"Abas-índice/guia detectadas: {', '.join(metas)} (são o MAPA do doc)."
           if metas else "Nenhuma aba-índice/guia detectada pelo nome — confira se há um dicionário de abas.")
    ]
    if truncated:
        notes.append(
            "Mesmo no teto ampliado a amostra foi cortada — abas de conteúdo no fim podem ter ficado "
            "de fora; use o INVENTÁRIO e o Guia para raciocinar sobre elas."
        )
    return text, truncated, notes


# ── Hint de fatiamento (por CÉLULAS e nº de sheets, não só linhas) ─────
def _chunking_hint(sheets_struct: list[dict[str, Any]]) -> str | None:
    """Sinaliza ao extrator quando o doc precisa ser fatiado pra não estourar
    a saída/contexto. Baseado em células (linhas×colunas) e nº de sheets."""
    big_sheets = [
        s for s in sheets_struct if s["totalRows"] * max(1, s["totalCols"]) > _XLSX_BIG_CELLS
    ]
    total_cells = sum(s["totalRows"] * max(1, s["totalCols"]) for s in sheets_struct)
    if not big_sheets and total_cells <= _XLSX_BIG_CELLS and len(sheets_struct) <= 3:
        return None
    partes = []
    if len(sheets_struct) > 3:
        partes.append(f"{len(sheets_struct)} sheets")
    for s in big_sheets:
        partes.append(f"'{s['name']}' {s['totalRows']}×{s['totalCols']}")
    detalhe = "; ".join(partes) or f"{total_cells} células no total"
    return (
        f"Planilha volumosa ({detalhe}) — extrair POR SHEET com ingerir_planilha (lê as células "
        "em código; não transcreva linha a linha). Confira a linha de totais (cauda) e registre em "
        "alerta como cada sheet foi estruturada."
    )


# ── XLSX ───────────────────────────────────────────────────────────
def _sample_xlsx(filename: str, data: bytes) -> DocSample:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    sheet_names = list(wb.sheetnames)
    # Modo workbook-motor: planilha multi-aba com índice/guia → inventário + meta-first + teto maior.
    workbook_mode = len(sheet_names) > _XLSX_WORKBOOK_SHEETS

    blocks: list[tuple[bool, str]] = []  # (is_meta, render) — ordem define o que sobrevive ao clip
    sheets_struct: list[dict[str, Any]] = []
    any_data = False

    for name in sheet_names:
        ws = wb[name]
        is_meta = workbook_mode and _is_meta_sheet(name)
        head_cap = _XLSX_META_HEAD_ROWS if is_meta else _XLSX_HEAD_ROWS
        head: list[tuple] = []
        tail: deque = deque(maxlen=_XLSX_TAIL_ROWS)
        mid: list[tuple[int, tuple]] = []  # (idx 1-based, row) · janela do MEIO
        max_cols = 0

        # Total robusto: max_row quando confiável, senão contagem (o mapper roda 1x/doc).
        total = ws.max_row if isinstance(ws.max_row, int) and ws.max_row > 0 else 0
        if total <= 0:
            total = sum(1 for _ in ws.iter_rows(values_only=True))

        # Janela do meio — fecha o ponto cego de blocos empilhados/cabeçalho-repetido
        # no miolo de planilha grande.
        mid_lo = mid_hi = -1
        if total > (head_cap + _XLSX_TAIL_ROWS + _XLSX_MID_ROWS + 4):
            half = total // 2
            mid_lo, mid_hi = half - _XLSX_MID_ROWS // 2, half + _XLSX_MID_ROWS // 2

        # Passada de captura · memória O(head + meio + tail), não O(linhas).
        idx = 0
        for row in ws.iter_rows(values_only=True):
            idx += 1
            max_cols = max(max_cols, len(row))
            if len(head) < head_cap + 1:  # +1 = linha de cabeçalho
                head.append(row)
            if mid_lo <= idx <= mid_hi:
                mid.append((idx, row))
            tail.append(row)

        total = max(total, idx)
        if total > 0:
            any_data = True

        blocks.append(
            (is_meta, _render_sheet(name, head, list(tail), total, max_cols, mid=mid,
                                    head_rows=head_cap, is_meta=is_meta))
        )
        sheets_struct.append(
            {
                "name": name,
                "totalRows": total,
                "totalCols": max_cols,  # largura REAL (não capada) — o extrator precisa saber
                "sampledCols": min(max_cols, _XLSX_MAX_COLS),  # quantas foram mostradas na amostra
                "truncatedCols": max_cols > _XLSX_MAX_COLS,
                "isMeta": is_meta,
            }
        )

    wb.close()

    text, truncated, notes = _assemble_xlsx(blocks, sheets_struct, workbook_mode)
    return DocSample(
        text=text,
        fmt="xlsx",
        total_units=len(sheet_names),
        readable=any_data,
        truncated=truncated,
        structure={
            "format": "xlsx",
            "totalUnits": len(sheet_names),
            "sheets": sheets_struct,
            "anomalies": [],
            "chunkingStrategy": _chunking_hint(sheets_struct),
            "sampledBy": "python-sampler@1.0.0",
        },
        notes=(notes if any_data else ["Planilha vazia ou sem linhas legíveis."]),
    )


def _render_sheet(
    name: str,
    head: list[tuple],
    tail: list[tuple],
    total: int,
    max_cols: int,
    mid: list[tuple[int, tuple]] | None = None,
    head_rows: int = _XLSX_HEAD_ROWS,
    is_meta: bool = False,
) -> str:
    cols = min(max_cols, _XLSX_MAX_COLS)
    col_note = f" (mostrando {cols} de {max_cols} colunas)" if max_cols > cols else ""
    tag = " · ABA-META (índice/guia — o MAPA do doc)" if is_meta else ""
    lines = [f"### Sheet \"{name}\" · {total} linhas × {max_cols} colunas{col_note}{tag}"]

    def fmt_row(idx: int, row: tuple) -> str:
        cells = []
        for v in row[:cols]:
            s = "" if v is None else str(v)
            s = s.replace("\n", " ").replace("\t", " ")
            if len(s) > _XLSX_CELL_CHARS:
                s = s[:_XLSX_CELL_CHARS] + "…"
            cells.append(s)
        return f"L{idx}: " + " | ".join(cells)

    # Cabeçalho + primeiras linhas
    shown_head = head[: head_rows + 1]
    for i, row in enumerate(shown_head, start=1):
        lines.append(fmt_row(i, row))

    head_count = len(shown_head)

    # Amostra do MEIO (se capturada e não sobrepõe head/tail) — fecha o ponto cego.
    if mid:
        mid_rows = [(idx, r) for (idx, r) in mid if idx > head_count and idx < total - len(tail) + 1]
        if mid_rows:
            lines.append(f"… ({mid_rows[0][0] - head_count - 1} linhas omitidas) …")
            lines.append("— amostra do MEIO —")
            for idx, row in mid_rows:
                lines.append(fmt_row(idx, row))

    # Gap + cauda (se houver linhas não mostradas no meio)
    if total > head_count + len(tail):
        omitted = total - head_count - len(tail)
        lines.append(f"… ({omitted} linhas omitidas) …")
        start_idx = total - len(tail) + 1
        for j, row in enumerate(tail):
            lines.append(fmt_row(start_idx + j, row))
    elif total > head_count:
        # Sobreposição pequena · só completa o que faltou sem duplicar.
        extra = tail[head_count - (total - len(tail)) :] if total > len(tail) else []
        start_idx = head_count + 1
        for j, row in enumerate(extra):
            lines.append(fmt_row(start_idx + j, row))

    return "\n".join(lines)


# ── XLS (binário legado) ───────────────────────────────────────────
def _sample_xls(filename: str, data: bytes) -> DocSample:
    import xlrd  # lê apenas .xls (xlrd 2.x dropou .xlsx)

    book = xlrd.open_workbook(file_contents=data)
    sheet_names = list(book.sheet_names())
    workbook_mode = len(sheet_names) > _XLSX_WORKBOOK_SHEETS

    blocks: list[tuple[bool, str]] = []
    sheets_struct: list[dict[str, Any]] = []
    any_data = False

    for name in sheet_names:
        sh = book.sheet_by_name(name)
        nrows, ncols = sh.nrows, sh.ncols
        if nrows > 0:
            any_data = True
        is_meta = workbook_mode and _is_meta_sheet(name)
        head_cap = _XLSX_META_HEAD_ROWS if is_meta else _XLSX_HEAD_ROWS
        # xlrd tem acesso aleatório → head, meio e tail direto por índice.
        head = [tuple(sh.row_values(r)) for r in range(min(nrows, head_cap + 1))]
        tail = [tuple(sh.row_values(r)) for r in range(max(0, nrows - _XLSX_TAIL_ROWS), nrows)]
        mid: list[tuple[int, tuple]] = []
        if nrows > (head_cap + _XLSX_TAIL_ROWS + _XLSX_MID_ROWS + 4):
            half = nrows // 2
            lo = half - _XLSX_MID_ROWS // 2
            mid = [(r + 1, tuple(sh.row_values(r))) for r in range(lo, lo + _XLSX_MID_ROWS)]
        blocks.append(
            (is_meta, _render_sheet(name, head, tail, nrows, ncols, mid=mid,
                                    head_rows=head_cap, is_meta=is_meta))
        )
        sheets_struct.append(
            {
                "name": name,
                "totalRows": nrows,
                "totalCols": ncols,  # largura REAL (não capada)
                "sampledCols": min(ncols, _XLSX_MAX_COLS),
                "truncatedCols": ncols > _XLSX_MAX_COLS,
                "isMeta": is_meta,
            }
        )

    text, truncated, notes = _assemble_xlsx(blocks, sheets_struct, workbook_mode)
    return DocSample(
        text=text,
        fmt="xls",
        total_units=len(sheet_names),
        readable=any_data,
        truncated=truncated,
        structure={
            "format": "xls",
            "totalUnits": len(sheet_names),
            "sheets": sheets_struct,
            "anomalies": [],
            "chunkingStrategy": _chunking_hint(sheets_struct),
            "sampledBy": "python-sampler@1.0.0",
        },
        notes=(notes if any_data else ["Planilha .xls vazia ou sem linhas legíveis."]),
    )


# ── DOCX ───────────────────────────────────────────────────────────
def _sample_docx(filename: str, data: bytes) -> DocSample:
    from docx import Document

    doc = Document(io.BytesIO(data))
    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    body = "\n".join(paras)

    table_parts: list[str] = []
    for ti, tbl in enumerate(doc.tables[:5], start=1):
        rows_txt = []
        for r in tbl.rows[:15]:
            cells = [c.text.replace("\n", " ").strip()[:_XLSX_CELL_CHARS] for c in r.cells]
            rows_txt.append(" | ".join(cells))
        if rows_txt:
            table_parts.append(f"#### Tabela {ti}\n" + "\n".join(rows_txt))

    n_paras = len(paras)
    n_tables = len(doc.tables)
    combined = body[:_DOCX_CHARS]
    if table_parts:
        combined += "\n\n" + "\n\n".join(table_parts)

    text, truncated = _clip(combined)
    readable = bool(body.strip()) or bool(table_parts)
    return DocSample(
        text=text or "(documento sem texto)",
        fmt="docx",
        total_units=n_tables,
        readable=readable,
        truncated=truncated,
        structure={
            "format": "docx",
            "totalUnits": n_tables,
            "paragraphs": n_paras,
            "tables": n_tables,
            "anomalies": [],
            "chunkingStrategy": None,
            "sampledBy": "python-sampler@1.0.0",
        },
        notes=([] if readable else ["DOCX sem texto extraível."]),
    )


# ── CSV / texto ────────────────────────────────────────────────────
def _decode(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _sample_text_lines(filename: str, data: bytes, fmt: str) -> DocSample:
    full = _decode(data)
    all_lines = full.splitlines()
    n = len(all_lines)
    shown = all_lines[:_CSV_LINES]
    body = "\n".join(shown)
    if n > _CSV_LINES:
        body += f"\n… ({n - _CSV_LINES} linhas restantes omitidas) …"
    text, truncated = _clip(body)
    return DocSample(
        text=text or "(arquivo vazio)",
        fmt=fmt,
        total_units=n,
        readable=bool(full.strip()),
        truncated=truncated or n > _CSV_LINES,
        structure={
            "format": fmt,
            "totalUnits": n,
            "anomalies": [],
            "chunkingStrategy": (f"CSV com {n} linhas — extrair em batches." if n > 500 else None),
            "sampledBy": "python-sampler@1.0.0",
        },
    )


def _sample_text_chars(filename: str, data: bytes, fmt: str) -> DocSample:
    full = _decode(data)
    text, truncated = _clip(full[:_TEXT_CHARS], hard=True)
    if len(full) > _TEXT_CHARS:
        truncated = True
    return DocSample(
        text=text or "(arquivo vazio)",
        fmt=fmt,
        total_units=0,
        readable=bool(full.strip()),
        truncated=truncated,
        structure={
            "format": fmt,
            "totalUnits": 0,
            "chars": len(full),
            "anomalies": [],
            "chunkingStrategy": None,
            "sampledBy": "python-sampler@1.0.0",
        },
    )


# ── helpers ────────────────────────────────────────────────────────
def _unsupported(filename: str, ext: str, why: str) -> DocSample:
    return DocSample(
        text=f"(arquivo '{filename}' · formato .{ext}) {why}",
        fmt=ext if ext in ("pdf", "xlsx", "docx", "csv", "md") else "unknown",
        total_units=0,
        readable=False,
        truncated=False,
        structure={"format": ext or "unknown", "totalUnits": 0, "unsupported": True},
        notes=[why],
    )


def _clip(text: str, hard: bool = False, cap: int | None = None) -> tuple[str, bool]:
    """Corta no teto de caracteres. Retorna (texto, truncado?). `cap` sobrescreve o teto global
    (usado no modo workbook, que tem um teto maior)."""
    from config import MAX_SAMPLE_CHARS

    limit = cap if cap is not None else MAX_SAMPLE_CHARS
    if len(text) <= limit:
        return text, False
    clipped = text[:limit]
    suffix = "" if hard else "\n\n… (amostra truncada no limite de contexto) …"
    return clipped + suffix, True
