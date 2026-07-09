"""Tools in-process MCP de LEITURA do documento (Fase 2).

O extrator lê o documento INTEIRO sob demanda por estas tools (≠ do mapper, que
só amostrava). Cada tool é uma closure sobre um `DocContext` (o doc aberto), e o
servidor é construído por-documento via `build_doc_tools_server(doc)`.

Texto: PDF texto, planilha (XLSX/XLS) por janelas, DOCX, CSV, MD.
Visão/OCR: `ler_pdf_pagina_imagem` (rasteriza p/ visão · PyMuPDF) e `ocr_pagina`
(rasteriza + tesseract). Degradam com aviso quando o binário/traineddata falta.
"""

from __future__ import annotations

import base64
import io
import re
import zipfile
from typing import Any
from xml.etree import ElementTree as ET

from claude_agent_sdk import create_sdk_mcp_server, tool

from config import OCR_LANG, PDF_RASTER_SCALE, XLSX_RANGE_MAX_ROWS

_CELL_CHARS = 200  # corta célula gigante na renderização


# ── Contexto do documento aberto ───────────────────────────────────────
class DocContext:
    def __init__(self, filename: str, data: bytes):
        self.filename = filename
        self.data = data
        self.ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        self._pdf = None
        self._fitz = None
        self._sheet_rows: dict[str, list[tuple]] = {}  # cache sheet → rows
        self._sheet_names: list[str] | None = None
        # {sheet: {(linha,coluna) 1-based}} de fórmulas SEM valor em cache (openpyxl
        # devolve None) → a ingestão sinaliza needs_review em vez de perder dado calado.
        self._uncached_formula: dict[str, set] = {}

    # PDF (pypdf · texto).
    def pdf(self):
        if self._pdf is None:
            from pypdf import PdfReader

            self._pdf = PdfReader(io.BytesIO(self.data))
        return self._pdf

    # PDF (PyMuPDF/fitz · rasterização p/ visão + OCR).
    def fitz(self):
        if self._fitz is None:
            import fitz  # PyMuPDF

            self._fitz = fitz.open(stream=self.data, filetype="pdf")
        return self._fitz

    def raster_page_png(self, page_1based: int, scale: float | None = None) -> bytes:
        """Rasteriza 1 página do PDF como PNG (1-based). Para visão/OCR."""
        import fitz  # PyMuPDF

        s = float(scale or PDF_RASTER_SCALE)
        doc = self.fitz()
        n = doc.page_count
        if page_1based < 1 or page_1based > n:
            raise ValueError(f"página {page_1based} fora do intervalo (PDF tem {n})")
        page = doc.load_page(page_1based - 1)
        pix = page.get_pixmap(matrix=fitz.Matrix(s, s))
        return pix.tobytes("png")

    def pdf_table_rows(self, page_1based: int) -> list[list[list]]:
        """Tabelas de TEXTO de uma página (PyMuPDF find_tables) — grade lida em código,
        determinística. Cada tabela = list[list[str]]. Vazio = página sem tabela de
        texto (provável imagem/escaneada → usar visão/OCR)."""
        d = self.fitz()
        n = d.page_count
        if page_1based < 1 or page_1based > n:
            raise ValueError(f"página {page_1based} fora do intervalo (PDF tem {n})")
        page = d.load_page(page_1based - 1)
        try:
            finder = page.find_tables()
        except Exception:  # noqa: BLE001 — versão sem find_tables ou página problemática
            return []
        out: list[list[list]] = []
        for t in getattr(finder, "tables", []) or []:
            try:
                rows = t.extract()
            except Exception:  # noqa: BLE001
                continue
            if rows:
                out.append(rows)
        return out

    # Planilha (xlsx/xlsm via openpyxl · xls via xlrd) · carrega linhas 1x e cacheia.
    # Trata mescladas (preenche o valor da âncora) e detecta fórmulas sem cache.
    def _ensure_sheets(self) -> None:
        if self._sheet_names is not None:
            return
        if self.ext in ("xlsx", "xlsm"):
            from openpyxl import load_workbook

            wb = load_workbook(io.BytesIO(self.data), read_only=True, data_only=True)
            self._sheet_names = list(wb.sheetnames)
            grids: dict[str, list[list]] = {}
            for name in self._sheet_names:
                ws = wb[name]
                # NÃO confiar no atributo `dimension` do arquivo: exports de ERP/geradores
                # mentem (ex.: 'A1:B3' numa sheet 10×5) e o read_only truncaria linhas e
                # colunas EM SILÊNCIO — re-deriva a extensão real varrendo o XML.
                if hasattr(ws, "reset_dimensions"):
                    ws.reset_dimensions()
                grids[name] = [list(r) for r in ws.iter_rows(values_only=True)]
            wb.close()
            # Um streaming por sheet: mescladas (read_only não expõe merged_cells) +
            # quais sheets têm fórmula. Propaga o valor da âncora das mescladas.
            merges, formula_sheets = _xlsx_scan(self.data)
            self._fill_merges(grids, merges)
            self._sheet_rows = grids  # listas MUTÁVEIS (temporário · permite preencher)
            # Fórmula sem valor em cache → None. Só paga a 2ª leitura cara nas sheets
            # que REALMENTE têm fórmula (export puro de dados pula isso). Depois CALCULA
            # as que faltam (pycel); o que sobrar (função não suportada) vai pra revisão.
            if formula_sheets:
                self._detect_uncached_formulas_xlsx(formula_sheets)
                if self._uncached_formula:
                    self._compute_uncached_formulas()
            self._sheet_rows = {n: [tuple(r) for r in g] for n, g in grids.items()}  # congela
        elif self.ext == "xls":
            import xlrd

            # formatting_info=True habilita merged_cells no xls — mas pode falhar em
            # alguns arquivos; cai pra sem-formatação (perde só as mescladas) sem quebrar.
            try:
                book = xlrd.open_workbook(file_contents=self.data, formatting_info=True)
            except Exception:  # noqa: BLE001
                book = xlrd.open_workbook(file_contents=self.data)
            self._sheet_names = list(book.sheet_names())
            grids = {}
            for name in self._sheet_names:
                sh = book.sheet_by_name(name)
                grid = []
                for r in range(sh.nrows):
                    vals = []
                    for c in range(sh.ncols):
                        cell = sh.cell(r, c)
                        ct = cell.ctype
                        if ct == xlrd.XL_CELL_DATE:
                            # xls guarda data como serial Excel — tipa de verdade (senão
                            # vira número sem sentido na ingestão). datemode = 1900/1904.
                            try:
                                vals.append(xlrd.xldate_as_datetime(cell.value, book.datemode))
                            except Exception:  # noqa: BLE001
                                vals.append(cell.value)
                        elif ct == xlrd.XL_CELL_EMPTY:
                            vals.append(None)
                        elif ct == xlrd.XL_CELL_BOOLEAN:
                            vals.append(bool(cell.value))
                        else:
                            vals.append(cell.value)
                    grid.append(vals)
                # mescladas do xls (xlrd dá [rlo,rhi) × [clo,chi), 0-based)
                ranges = []
                for rlo, rhi, clo, chi in getattr(sh, "merged_cells", []) or []:
                    ranges.append((rlo + 1, clo + 1, rhi, chi))  # → (min_r,min_c,max_r,max_c) 1-based incl.
                grids[name] = grid
                if ranges:
                    self._fill_merges({name: grid}, {name: ranges})
            self._sheet_rows = {n: [tuple(r) for r in g] for n, g in grids.items()}
        else:
            self._sheet_names = []

    @staticmethod
    def _fill_merges(grids: dict, ranges_by_sheet: dict) -> None:
        """Propaga o valor da âncora de cada região mesclada p/ as células não-âncora
        (que vêm vazias) — senão a categoria/label mesclada some nas linhas de baixo."""
        for name, ranges in (ranges_by_sheet or {}).items():
            grid = grids.get(name)
            if not grid:
                continue
            for min_r, min_c, max_r, max_c in ranges:
                if min_r - 1 >= len(grid):
                    continue
                arow = grid[min_r - 1]
                anchor = arow[min_c - 1] if min_c - 1 < len(arow) else None
                if anchor is None:
                    continue
                for rr in range(min_r, max_r + 1):
                    if rr - 1 >= len(grid):
                        break
                    row = grid[rr - 1]
                    for cc in range(min_c, max_c + 1):
                        while len(row) < cc:
                            row.append(None)
                        if row[cc - 1] is None:
                            row[cc - 1] = anchor

    def _detect_uncached_formulas_xlsx(self, only_sheets=None) -> None:  # noqa: ANN001
        """Marca células-fórmula cujo valor NÃO está em cache (openpyxl→None). A
        ingestão usa isso pra mandar pra needs_review em vez de anexar dado faltando.
        `only_sheets` limita às sheets que têm fórmula (evita custo nas outras)."""
        try:
            from openpyxl import load_workbook

            wbf = load_workbook(io.BytesIO(self.data), read_only=True, data_only=False)
        except Exception:  # noqa: BLE001
            return
        try:
            for name in self._sheet_names or []:
                if only_sheets is not None and name not in only_sheets:
                    continue
                try:
                    wsf = wbf[name]
                except Exception:  # noqa: BLE001
                    continue
                vrows = self._sheet_rows.get(name) or []
                uncached = set()
                for r, row in enumerate(wsf.iter_rows(), start=1):
                    for c, cell in enumerate(row, start=1):
                        v = cell.value
                        if isinstance(v, str) and v.startswith("="):
                            cached = (
                                vrows[r - 1][c - 1]
                                if r - 1 < len(vrows) and c - 1 < len(vrows[r - 1])
                                else None
                            )
                            if cached is None:
                                uncached.add((r, c))
                if uncached:
                    self._uncached_formula[name] = uncached
        finally:
            try:
                wbf.close()
            except Exception:  # noqa: BLE001
                pass

    def _compute_uncached_formulas(self) -> None:
        """CALCULA (pycel) as fórmulas sem valor em cache e preenche o grid. O que
        a engine não resolver (função não suportada / erro) FICA em `_uncached_formula`
        → vira needs_review. Mutável: roda antes de congelar o grid em tuplas."""
        if not self._uncached_formula:
            return
        try:
            import logging
            import numbers

            from openpyxl.utils import get_column_letter
            from pycel import ExcelCompiler

            # pycel loga traceback por célula em função não suportada → silencia o spam
            # (a célula não resolvida já é sinalizada pela nossa flag de revisão).
            logging.getLogger("pycel").setLevel(logging.CRITICAL)
        except Exception:  # noqa: BLE001 — pycel ausente → mantém fail-loud (revisão)
            return
        import os
        import tempfile

        fd, path = tempfile.mkstemp(suffix=".xlsx")
        try:
            with os.fdopen(fd, "wb") as fh:
                fh.write(self.data)
            try:
                xl = ExcelCompiler(filename=path)
            except Exception:  # noqa: BLE001
                return
            for name in list(self._uncached_formula.keys()):
                grid = self._sheet_rows.get(name)
                cells = self._uncached_formula.get(name) or set()
                if not grid:
                    continue
                ainda: set = set()
                for r, c in cells:
                    addr = f"'{name}'!{get_column_letter(c)}{r}"
                    try:
                        val = xl.evaluate(addr)
                    except Exception:  # noqa: BLE001
                        val = None
                    # erro Excel ('#DIV/0!'…) ou vazio → não resolvido (mantém flag)
                    if val is None or (isinstance(val, str) and val.startswith("#")):
                        ainda.add((r, c))
                        continue
                    # normaliza tipos numpy do pycel p/ tipos nativos
                    if isinstance(val, bool):
                        pass
                    elif isinstance(val, numbers.Integral):
                        val = int(val)
                    elif isinstance(val, numbers.Real):
                        val = float(val)
                    if r - 1 < len(grid):
                        row = grid[r - 1]
                        while len(row) < c:
                            row.append(None)
                        if row[c - 1] is None:
                            row[c - 1] = val
                if ainda:
                    self._uncached_formula[name] = ainda
                else:
                    self._uncached_formula.pop(name, None)
        finally:
            try:
                os.remove(path)
            except Exception:  # noqa: BLE001
                pass

    def uncached_formula_cells(self, sheet: str, de: int, ate: int) -> list[str]:
        """Refs (ex.: 'D2') das fórmulas sem cache (NÃO resolvidas) dentro de [de, ate]."""
        self._ensure_sheets()
        cells = self._uncached_formula.get(sheet) or set()
        from openpyxl.utils import get_column_letter

        refs = [f"{get_column_letter(c)}{r}" for (r, c) in cells if de <= r <= ate]
        return sorted(refs)

    def sheet_names(self) -> list[str]:
        self._ensure_sheets()
        return self._sheet_names or []

    def sheet_rows(self, name: str) -> list[tuple]:
        self._ensure_sheets()
        return self._sheet_rows.get(name, [])

    def close(self) -> None:
        self._pdf = None
        if self._fitz is not None:
            try:
                self._fitz.close()
            except Exception:  # noqa: BLE001
                pass
        self._fitz = None
        self._sheet_rows = {}
        self._sheet_names = None
        self._uncached_formula = {}


# ── helpers de resposta MCP ────────────────────────────────────────────
def _txt(s: str) -> dict[str, Any]:
    return {"content": [{"type": "text", "text": s}]}


def _err(s: str) -> dict[str, Any]:
    return {"content": [{"type": "text", "text": s}], "is_error": True}


def _img(png: bytes, note: str) -> dict[str, Any]:
    b64 = base64.b64encode(png).decode("ascii")
    return {
        "content": [
            {"type": "text", "text": note},
            {"type": "image", "data": b64, "mimeType": "image/png"},
        ]
    }


def _cell(v: Any) -> str:
    if v is None:
        return ""
    s = str(v).replace("\n", " ").replace("\t", " ")
    return s[:_CELL_CHARS] + "…" if len(s) > _CELL_CHARS else s


def _render_rows(rows: list[tuple], start_idx: int) -> str:
    out = []
    for i, row in enumerate(rows):
        cells = " | ".join(_cell(v) for v in row)
        out.append(f"L{start_idx + i}: {cells}")
    return "\n".join(out)


TOOL_NAMES = [
    "dimensoes",
    "ler_pdf_paginas",
    "listar_tabelas_pdf",
    "ler_pdf_pagina_imagem",
    "ocr_pagina",
    "ler_planilha",
    "ler_docx",
    "ler_csv",
    "ler_markdown",
]


def build_doc_tools_server(doc: DocContext):
    """Cria o servidor MCP in-process com as tools de leitura amarradas a `doc`."""

    @tool("dimensoes", "Visão geral do documento (formato, nº de páginas/sheets/linhas). Chame PRIMEIRO.", {})
    async def dimensoes(args):  # noqa: ANN001
        ext = doc.ext
        if ext == "pdf":
            try:
                n = len(doc.pdf().pages)
            except Exception as e:  # noqa: BLE001
                return _err(f"PDF ilegível: {type(e).__name__}: {e}")
            return _txt(
                f"PDF · {n} página(s). Texto corrido: ler_pdf_paginas(de, ate) (≤10 por chamada). "
                "Para TABELAS (ex.: BM): chame listar_tabelas_pdf(pagina) PRIMEIRO — se achar tabela de texto, "
                "ingira com ingerir_tabela_pdf (grade lida em código · rápido e fiel). Só use ler_pdf_pagina_imagem "
                "(visão) como FALLBACK quando não houver tabela de texto (página escaneada). Sem texto algum: ocr_pagina(pagina)."
            )
        if ext in ("xlsx", "xlsm", "xls"):
            try:
                names = doc.sheet_names()
            except Exception as e:  # noqa: BLE001
                return _err(f"Planilha ilegível: {type(e).__name__}: {e}")
            linhas = []
            for nm in names:
                rows = doc.sheet_rows(nm)
                cols = max((len(r) for r in rows[:50]), default=0)
                linhas.append(f"  · '{nm}': {len(rows)} linhas × ~{cols} colunas")
            return _txt(
                f"Planilha · {len(names)} sheet(s):\n" + "\n".join(linhas) +
                f"\nFLUXO RÁPIDO: use ler_planilha(sheet, de, ate) (janela ≤ {XLSX_RANGE_MAX_ROWS}) só pra AMOSTRAR "
                "o cabeçalho/início e a cauda (totais). Depois ingira a tabela inteira de uma vez com "
                "ingerir_planilha(sheet, secao_id, titulo, linha_cabecalho, de, ate) — ela lê as células em código "
                "(rápido e fiel). NÃO transcreva linha a linha com anexar_linhas."
            )
        if ext == "docx":
            return _txt("DOCX · use ler_docx() (texto + tabelas).")
        if ext == "csv":
            return _txt("CSV · use ler_csv().")
        if ext in ("md", "markdown", "txt"):
            return _txt("Texto · use ler_markdown().")
        return _err(
            f"Formato .{ext} sem leitor direto aqui. Registre em alertas_extracao que precisa conversão "
            "(ex.: .mpp → exportar XML; .doc → .docx)."
        )

    @tool("ler_pdf_paginas", "Lê o TEXTO de um intervalo de páginas do PDF (máx 10 por chamada).", {"de": int, "ate": int})
    async def ler_pdf_paginas(args):  # noqa: ANN001
        if doc.ext != "pdf":
            return _err("Documento não é PDF.")
        try:
            reader = doc.pdf()
            total = len(reader.pages)
        except Exception as e:  # noqa: BLE001
            return _err(f"PDF ilegível: {e}")
        de = max(1, int(args.get("de", 1)))
        ate = min(total, int(args.get("ate", de)))
        if ate - de + 1 > 10:
            ate = de + 9
        parts = []
        for p in range(de, ate + 1):
            try:
                txt = (reader.pages[p - 1].extract_text() or "").strip()
            except Exception:  # noqa: BLE001
                txt = ""
            parts.append(f"=== Página {p}/{total} ===\n{txt or '(sem texto extraível · página pode ser imagem/escaneada)'}")
        return _txt("\n\n".join(parts))

    @tool(
        "listar_tabelas_pdf",
        "Detecta TABELAS DE TEXTO numa página de PDF (rápido, determinístico · PyMuPDF). Chame ANTES da visão: se achar tabela, ingira com ingerir_tabela_pdf (lê a grade em código — fiel e barato). Vazio = página é imagem/escaneada → aí sim use ler_pdf_pagina_imagem/ocr_pagina.",
        {"pagina": int},
    )
    async def listar_tabelas_pdf(args):  # noqa: ANN001
        if doc.ext != "pdf":
            return _err("Documento não é PDF.")
        p = max(1, int(args.get("pagina", 1)))
        try:
            tables = doc.pdf_table_rows(p)
        except Exception as e:  # noqa: BLE001
            return _err(str(e))
        if not tables:
            return _txt(
                f"Página {p}: nenhuma tabela de TEXTO detectada (pode ser imagem/escaneada). "
                "Use ler_pdf_pagina_imagem (visão) ou ocr_pagina."
            )
        lines = [f"Página {p}: {len(tables)} tabela(s) de texto detectada(s):"]
        for i, t in enumerate(tables):
            ncols = max((len(r) for r in t[:8]), default=0)
            lines.append(
                f"  #{i}: {len(t)} linhas × ~{ncols} colunas · PRÉVIA das 1ªs linhas "
                "(ache a linha do cabeçalho REAL — em BM/Medição ela costuma vir DEPOIS "
                "de linhas de título/metadados como Obra:/Período:/Contrato:/totais de página):"
            )
            for ri, row in enumerate(t[:7], start=1):
                cells = " | ".join(_cell(c) for c in row[:ncols])
                lines.append(f"      L{ri}: {cells[:200]}")
        lines.append(
            "Ingira com ingerir_tabela_pdf(pagina, indice_tabela, secao_id, titulo, linha_cabecalho=<nº L da linha do cabeçalho REAL>). "
            "Linhas de título/metadados ACIMA do cabeçalho são ignoradas automaticamente. "
            "Tabela que cruza páginas: MESMA secao_id; nas páginas seguintes o cabeçalho se repete — aponte linha_cabecalho pra ele de novo (a auditoria sinaliza se sobrar cabeçalho como dado)."
        )
        return _txt("\n".join(lines))

    @tool(
        "ler_pdf_pagina_imagem",
        "Rasteriza UMA página do PDF e te devolve a IMAGEM (visão). FALLBACK: use quando listar_tabelas_pdf não achar tabela de texto (página escaneada/imagem) ou a grade vier ruim. Leia os números direto da imagem.",
        {"pagina": int},
    )
    async def ler_pdf_pagina_imagem(args):  # noqa: ANN001
        if doc.ext != "pdf":
            return _err("Documento não é PDF.")
        try:
            total = len(doc.pdf().pages)
        except Exception:  # noqa: BLE001
            total = None
        p = max(1, int(args.get("pagina", 1)))
        if total and p > total:
            return _err(f"Página {p} fora do intervalo (PDF tem {total}).")
        try:
            png = doc.raster_page_png(p)
        except Exception as e:  # noqa: BLE001
            return _err(f"Falha ao rasterizar página {p}: {type(e).__name__}: {e}. Use ler_pdf_paginas (texto) e registre em alertas_extracao.")
        return _img(png, f"Página {p}{f'/{total}' if total else ''} (imagem · leia a tabela diretamente da imagem).")

    @tool(
        "ocr_pagina",
        "OCR de UMA página escaneada/sem texto extraível (rasteriza + tesseract). Use quando ler_pdf_paginas retorna vazio. Precisa do binário tesseract no host.",
        {"pagina": int},
    )
    async def ocr_pagina(args):  # noqa: ANN001
        if doc.ext != "pdf":
            return _err("Documento não é PDF.")
        p = max(1, int(args.get("pagina", 1)))
        try:
            png = doc.raster_page_png(p)
        except Exception as e:  # noqa: BLE001
            return _err(f"Falha ao rasterizar página {p}: {type(e).__name__}: {e}")
        try:
            import pytesseract
            from PIL import Image

            text = pytesseract.image_to_string(Image.open(io.BytesIO(png)), lang=OCR_LANG)
        except Exception as e:  # noqa: BLE001 — tesseract/traineddata ausente no host
            return _err(
                f"OCR indisponível ({type(e).__name__}: {e}). Tente ler_pdf_pagina_imagem(pagina) (visão) "
                "para ler a página; se nada funcionar, registre em alertas_extracao que precisa de OCR (tesseract+'por')."
            )
        text = (text or "").strip()
        return _txt(f"OCR pág {p} (lang={OCR_LANG}):\n{text or '(OCR não extraiu texto desta página)'}")

    @tool("ler_planilha", "Lê um intervalo de linhas de uma sheet (de..ate, 1-based). Janela ≤ limite.", {"sheet": str, "de": int, "ate": int})
    async def ler_planilha(args):  # noqa: ANN001
        if doc.ext not in ("xlsx", "xlsm", "xls"):
            return _err("Documento não é planilha.")
        names = doc.sheet_names()
        sheet = args.get("sheet") or (names[0] if names else "")
        if sheet not in names:
            return _err(f"Sheet '{sheet}' não existe. Sheets: {names}")
        rows = doc.sheet_rows(sheet)
        total = len(rows)
        de = max(1, int(args.get("de", 1)))
        ate = int(args.get("ate", de))
        if ate < de:
            ate = de
        if ate - de + 1 > XLSX_RANGE_MAX_ROWS:
            ate = de + XLSX_RANGE_MAX_ROWS - 1
        ate = min(ate, total)
        window = rows[de - 1 : ate]
        body = _render_rows(window, de)
        more = (
            f"\n… (sheet tem {total} linhas; você leu {de}-{ate}). Para CAPTURAR a tabela, NÃO leia tudo aqui — "
            f"chame ingerir_planilha(sheet='{sheet}', linha_cabecalho=…, de=…, ate=…) (lê as células em código)."
            if ate < total
            else ""
        )
        return _txt(f"Sheet '{sheet}' · linhas {de}-{ate} de {total} (AMOSTRA):\n{body}{more}")

    @tool("ler_docx", "Lê todo o texto + tabelas de um DOCX.", {})
    async def ler_docx(args):  # noqa: ANN001
        if doc.ext != "docx":
            return _err("Documento não é DOCX.")
        try:
            from docx import Document

            d = Document(io.BytesIO(doc.data))
        except Exception as e:  # noqa: BLE001
            return _err(f"DOCX ilegível: {e}")
        paras = [p.text for p in d.paragraphs if p.text and p.text.strip()]
        out = ["\n".join(paras)]
        for ti, tbl in enumerate(d.tables, start=1):
            # NÃO trunca a célula na EXTRAÇÃO (o corte de 200 chars é só p/ amostra do mapper).
            rows = [" | ".join((c.text or "").replace("\n", " ").replace("\t", " ") for c in r.cells) for r in tbl.rows]
            out.append(f"#### Tabela {ti}\n" + "\n".join(rows))
        return _txt("\n\n".join(out) or "(documento vazio)")

    @tool("ler_csv", "Lê o conteúdo de um CSV.", {})
    async def ler_csv(args):  # noqa: ANN001
        if doc.ext != "csv":
            return _err("Documento não é CSV.")
        return _txt(_with_trunc_marker(_decode(doc.data), "CSV"))

    @tool("ler_markdown", "Lê o conteúdo de um arquivo de texto/markdown.", {})
    async def ler_markdown(args):  # noqa: ANN001
        if doc.ext not in ("md", "markdown", "txt"):
            return _err("Documento não é texto/markdown.")
        return _txt(_with_trunc_marker(_decode(doc.data), "Texto"))

    return create_sdk_mcp_server(
        "doctools",
        "1.0.0",
        tools=[
            dimensoes,
            ler_pdf_paginas,
            listar_tabelas_pdf,
            ler_pdf_pagina_imagem,
            ocr_pagina,
            ler_planilha,
            ler_docx,
            ler_csv,
            ler_markdown,
        ],
    )


_TEXT_LIMIT = 120_000


def _with_trunc_marker(text: str, kind: str) -> str:
    """Corta texto grande mas DEIXA VISÍVEL que cortou — pra o modelo registrar um
    anexar_alerta em vez de perder dados em silêncio."""
    if len(text) <= _TEXT_LIMIT:
        return text
    return (
        text[:_TEXT_LIMIT]
        + f"\n\n[⚠ {kind} TRUNCADO em {_TEXT_LIMIT} de {len(text)} chars — você NÃO viu o resto. "
        "Registre com anexar_alerta que o documento não foi lido por completo e quais linhas/seções podem faltar.]"
    )


def _decode(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _local(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


_MERGE_REF_RE = re.compile(rb'ref="([^"]+)"')
_FORMULA_RE = re.compile(rb"<f[ >/]")  # <f> | <f ...> | <f/> → célula com fórmula


def _scan_merge_refs(fh, chunk_size: int = 1 << 16):
    """Stream do XML da sheet (memória CONSTANTE) que devolve `(refs, has_formula)`:
    os refs de mescladas ('A2:A3') e se a sheet tem QUALQUER fórmula. O bloco
    <mergeCells> fica depois do sheetData (o grosso), então descartamos os chunks
    até achá-lo — e de quebra marcamos formula no mesmo passo (de graça), pra evitar
    a 2ª leitura cara só-pra-fórmula quando a planilha não tem nenhuma."""
    OPEN, CLOSE = b"<mergeCells", b"</mergeCells>"
    tail = b""  # fim do chunk anterior p/ casar padrão que cruza fronteira
    in_block = False
    block = b""
    refs: list[str] = []
    has_formula = False
    while True:
        chunk = fh.read(chunk_size)
        if not chunk:
            break
        buf = tail + chunk
        if not in_block:
            if not has_formula and _FORMULA_RE.search(buf):
                has_formula = True
            i = buf.find(OPEN)
            if i < 0:
                tail = buf[-(max(len(OPEN), 3) - 1):]  # pode cruzar fronteira
                continue
            in_block = True
            buf = buf[i:]
        j = buf.find(CLOSE)
        if j >= 0:
            block += buf[:j]
            break
        block += buf
        if len(block) > len(CLOSE):
            keep = block[-(len(CLOSE) - 1):]
            body = block[:-(len(CLOSE) - 1)]
            for m in _MERGE_REF_RE.finditer(body):
                refs.append(m.group(1).decode())
            block = keep
        tail = b""
    for m in _MERGE_REF_RE.finditer(block):
        refs.append(m.group(1).decode())
    return refs, has_formula


def _xlsx_scan(data: bytes) -> tuple[dict[str, list[tuple]], set[str]]:
    """Lê do XML do xlsx, em UM streaming por sheet (memória constante): os ranges de
    mescladas (read_only do openpyxl não os expõe) E quais sheets têm fórmula.
    Retorna ({sheet: [(min_r,min_c,max_r,max_c)] 1-based incl.}, {sheets com fórmula})."""
    out: dict[str, list[tuple]] = {}
    formula_sheets: set[str] = set()
    try:
        from openpyxl.utils import range_boundaries

        zf = zipfile.ZipFile(io.BytesIO(data))
    except Exception:  # noqa: BLE001
        return out, formula_sheets
    try:
        # workbook.xml: nome da sheet → r:id
        sheets: list[tuple[str, str]] = []
        try:
            wbx = ET.fromstring(zf.read("xl/workbook.xml"))
            for el in wbx.iter():
                if _local(el.tag) == "sheet":
                    name = el.attrib.get("name")
                    rid = next((v for k, v in el.attrib.items() if _local(k) == "id"), None)
                    if name and rid:
                        sheets.append((name, rid))
        except Exception:  # noqa: BLE001
            return out
        # rels: r:id → arquivo da sheet
        rid_target: dict[str, str] = {}
        try:
            relx = ET.fromstring(zf.read("xl/_rels/workbook.xml.rels"))
            for el in relx.iter():
                if _local(el.tag) == "Relationship":
                    rid_target[el.attrib.get("Id")] = el.attrib.get("Target")
        except Exception:  # noqa: BLE001
            pass
        import posixpath

        names_in_zip = set(zf.namelist())
        for name, rid in sheets:
            target = rid_target.get(rid)
            if not target:
                continue
            # Target pode ser absoluto ("/xl/worksheets/sheet1.xml") ou relativo a xl/.
            if target.startswith("/"):
                path = target.lstrip("/")
            else:
                path = posixpath.normpath(posixpath.join("xl", target))
            if path not in names_in_zip:
                continue
            ranges: list[tuple] = []
            try:
                with zf.open(path) as fh:
                    refs, has_formula = _scan_merge_refs(fh)
                if has_formula:
                    formula_sheets.add(name)
                for ref in refs:
                    if ":" in ref:
                        try:
                            min_c, min_r, max_c, max_r = range_boundaries(ref)
                            ranges.append((min_r, min_c, max_r, max_c))
                        except Exception:  # noqa: BLE001
                            pass
            except Exception:  # noqa: BLE001
                continue
            if ranges:
                out[name] = ranges
    finally:
        try:
            zf.close()
        except Exception:  # noqa: BLE001
            pass
    return out, formula_sheets
